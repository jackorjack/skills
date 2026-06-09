#!/usr/bin/env python3
"""docx2md — 将 .docx 文件或目录批量转换为 Markdown。

用法:
    python3 docx2md.py                          # 转换默认 wiki 目录下所有 docx
    python3 docx2md.py <目录>                    # 转换指定目录下所有 docx
    python3 docx2md.py <input.docx>             # 转换单个文件（同目录输出）
    python3 docx2md.py <input.docx> <out.md>    # 转换到指定路径
    python3 docx2md.py <input.docx> -           # 输出到 stdout

参数:
    --force     覆盖已存在的 md 文件（默认跳过）

依赖:
    pip install python-docx
"""

import sys
import os
import argparse
from docx import Document
from pathlib import Path

DEFAULT_DIR = os.path.expanduser("~/.openclaw/workspace/wiki")


def convert_docx_to_markdown(docx_path: str) -> str:
    """将 docx 文件转为 Markdown 字符串。"""
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        style_name = para.style.name if para.style else ''

        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
            except (IndexError, ValueError):
                level = 1
            level = min(level, 6)
            lines.append('#' * level + ' ' + text)

        elif style_name.startswith('List Bullet'):
            indent = style_name.count(' ')
            prefix = '  ' * (indent // 2) + '- '
            lines.append(prefix + text)

        elif style_name.startswith('List Number'):
            indent = style_name.count(' ')
            prefix = '  ' * (indent // 2) + '1. '
            lines.append(prefix + text)

        else:
            inline_parts = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold:
                    t = f'**{t}**'
                if run.italic:
                    t = f'*{t}*'
                inline_parts.append(t)
            lines.append(''.join(inline_parts) if inline_parts else text)

    for table in doc.tables:
        lines.append('')
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                lines.append('|' + '|'.join([' --- '] * len(cells)) + '|')
        lines.append('')

    return '\n'.join(lines)


def convert_one(docx_path: str, out_path: str | None, force: bool) -> tuple[int, int]:
    """转换单个文件。返回 (转换数, 跳过数)。"""
    md_path = out_path or os.path.splitext(docx_path)[0] + '.md'

    if os.path.isfile(md_path) and not force:
        return 0, 1

    md_text = convert_docx_to_markdown(docx_path)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_text)

    size = os.path.getsize(md_path)
    print(f'  ✅ {os.path.basename(md_path)} ({len(md_text.splitlines())} 行, {size:,} 字节)')
    return 1, 0


def convert_dir(dir_path: str, force: bool):
    """批量转换目录下所有 docx。"""
    if not os.path.isdir(dir_path):
        print(f'目录不存在: {dir_path}', file=sys.stderr)
        sys.exit(1)

    docx_files = sorted(Path(dir_path).glob('*.docx'))
    if not docx_files:
        print(f'目录下无 docx 文件: {dir_path}')
        return

    print(f'📂 {dir_path}  ({len(docx_files)} 个 docx)')

    converted = 0
    skipped = 0
    for docx in docx_files:
        c, s = convert_one(str(docx), None, force)
        converted += c
        skipped += s

    if skipped:
        print(f'\n📊 转换 {converted} 个, 跳过 {skipped} 个（md 已存在）')
    else:
        print(f'\n📊 转换 {converted} 个')


def main():
    parser = argparse.ArgumentParser(description='docx → Markdown 转换工具')
    parser.add_argument('input', nargs='?', default=DEFAULT_DIR,
                        help=f'文件或目录路径（默认: {DEFAULT_DIR}）')
    parser.add_argument('output', nargs='?', default=None,
                        help='输出文件路径（仅单文件模式；- 输出到 stdout）')
    parser.add_argument('--force', action='store_true',
                        help='强制覆盖已存在的 md 文件')
    args = parser.parse_args()

    input_path = os.path.expanduser(args.input)

    if input_path == '-':
        print('stdin 模式暂不支持', file=sys.stderr)
        sys.exit(1)

    # 目录模式
    if os.path.isdir(input_path):
        convert_dir(input_path, args.force)
        return

    # 单文件模式
    if not os.path.isfile(input_path):
        print(f'文件不存在: {input_path}', file=sys.stderr)
        sys.exit(1)

    if args.output == '-':
        print(convert_docx_to_markdown(input_path))
        return

    c, s = convert_one(input_path, args.output, args.force)
    if s:
        print(f'  ⏭ 跳过（md 已存在，用 --force 强制覆盖）')


if __name__ == '__main__':
    main()
